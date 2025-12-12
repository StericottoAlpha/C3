"""
ドキュメント解析サービス
PDF、Word、Markdownファイルからテキストを抽出
"""
import logging
from pathlib import Path
from typing import Optional

import PyPDF2
import docx
import markdown

logger = logging.getLogger(__name__)


class DocumentParser:
    """
    ドキュメント解析クラス
    様々なファイル形式からテキストを抽出
    """

    @classmethod
    def extract_text(cls, file_path: str, file_type: str) -> Optional[str]:
        """
        ファイル種別に応じてテキスト抽出

        Args:
            file_path: ファイルパス
            file_type: ファイル種別 ('pdf', 'docx', 'md')

        Returns:
            抽出されたテキスト（失敗時はNone）
        """
        try:
            if file_type == 'pdf':
                return cls.extract_from_pdf(file_path)
            elif file_type == 'docx':
                return cls.extract_from_docx(file_path)
            elif file_type == 'md':
                return cls.extract_from_markdown(file_path)
            else:
                logger.error(f"Unsupported file type: {file_type}")
                return None
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return None

    @classmethod
    def extract_from_pdf(cls, file_path: str) -> str:
        """
        PDFからテキスト抽出

        Args:
            file_path: PDFファイルパス

        Returns:
            抽出されたテキスト

        Raises:
            FileNotFoundError: ファイルが見つからない場合
            Exception: PDF解析エラー
        """
        if not Path(file_path).exists():
            raise FileNotFoundError(f"PDF file not found: {file_path}")

        text_content = []

        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            total_pages = len(pdf_reader.pages)

            for page_num in range(total_pages):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()

                if page_text:
                    # ページ番号情報を付加
                    text_content.append(f"[Page {page_num + 1}/{total_pages}]\n{page_text}")

        return "\n\n".join(text_content)

    @classmethod
    def extract_from_docx(cls, file_path: str) -> str:
        """
        Wordドキュメントからテキスト抽出

        Args:
            file_path: Wordファイルパス

        Returns:
            抽出されたテキスト

        Raises:
            FileNotFoundError: ファイルが見つからない場合
            Exception: Word解析エラー
        """
        if not Path(file_path).exists():
            raise FileNotFoundError(f"Word file not found: {file_path}")

        doc = docx.Document(file_path)
        text_content = []

        # パラグラフからテキスト抽出
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)

        # テーブルからテキスト抽出
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    text_content.append(row_text)

        return "\n\n".join(text_content)

    @classmethod
    def extract_from_markdown(cls, file_path: str) -> str:
        """
        Markdownファイルからテキスト抽出

        Args:
            file_path: Markdownファイルパス

        Returns:
            抽出されたテキスト（プレーンテキスト形式）

        Raises:
            FileNotFoundError: ファイルが見つからない場合
            Exception: Markdown解析エラー
        """
        if not Path(file_path).exists():
            raise FileNotFoundError(f"Markdown file not found: {file_path}")

        with open(file_path, 'r', encoding='utf-8') as file:
            md_content = file.read()

        # Markdownをプレーンテキストに変換
        # （実際のマニュアルはMarkdown構造も重要なので、そのまま返す）
        return md_content

    @classmethod
    def extract_metadata_from_pdf(cls, file_path: str) -> dict:
        """
        PDFからメタデータ抽出

        Args:
            file_path: PDFファイルパス

        Returns:
            メタデータ辞書
        """
        metadata = {
            'title': None,
            'author': None,
            'subject': None,
            'creator': None,
            'producer': None,
            'creation_date': None,
            'total_pages': 0,
        }

        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata['total_pages'] = len(pdf_reader.pages)

                if pdf_reader.metadata:
                    pdf_meta = pdf_reader.metadata
                    metadata['title'] = pdf_meta.get('/Title')
                    metadata['author'] = pdf_meta.get('/Author')
                    metadata['subject'] = pdf_meta.get('/Subject')
                    metadata['creator'] = pdf_meta.get('/Creator')
                    metadata['producer'] = pdf_meta.get('/Producer')
                    metadata['creation_date'] = pdf_meta.get('/CreationDate')

        except Exception as e:
            logger.warning(f"Failed to extract PDF metadata: {e}")

        return metadata
